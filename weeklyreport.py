import os
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import matplotlib.pyplot as plt
import io
from docx.text.paragraph import Paragraph

# Default statuses filter
focus_statuses = ["To Do", "In Progress", "QA"]

# ----------------------------------------
# Funciones auxiliares para inserciones en Word
# ----------------------------------------
def insert_paragraph_after(paragraph, text=None):
    """
    Inserta un nuevo párrafo justo después de `paragraph` y devuelve el objeto Paragraph creado.
    """
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para

def insert_table_after(doc, paragraph, rows, cols):
    """
    Crea una tabla de tamaño rows x cols al final del documento y luego la mueve
    para que quede justo después de `paragraph`. Devuelve la tabla insertada.
    """
    tbl = doc.add_table(rows=rows, cols=cols)
    body = doc._body._body
    body.remove(tbl._tbl)
    idx = list(body).index(paragraph._p)
    body.insert(idx + 1, tbl._tbl)
    return tbl

def set_column_width(table, column_index, width_inches):
    """
    Sets the width of a specific column in a table
    
    Args:
        table: The table to modify
        column_index: Zero-based index of the column to adjust
        width_inches: Width in inches
    """
    from docx.shared import Inches
    from docx.oxml import OxmlElement
    
    for row in table.rows:
        # Convert to twips (twentieth of a point) - the internal unit Word uses
        row.cells[column_index].width = Inches(width_inches)
        
        # Also set the width in the cell properties
        cell = row.cells[column_index]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        # Don't include namespace prefix in attribute names
        tcW.set('w', str(int(Inches(width_inches).twips)))
        tcW.set('type', 'dxa')
        
        # Remove existing width element if it exists
        # Fix: Use a simpler approach to find and remove tcW elements
        for child in tcPr.iterchildren():
            if child.tag.endswith('tcW'):
                tcPr.remove(child)
        
        tcPr.append(tcW)

def style_table_like_image(table):
    """
    Aplica el estilo de la imagen de muestra a la tabla:
    - Cabecera azul claro
    - Texto en negrita y centrado en la cabecera
    - Bordes en todas las celdas
    """
    # Import required modules
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Apply light blue header style (first row)
    header_row = table.rows[0]
    for cell in header_row.cells:
        # Set light blue background (#B8CCE4) for header
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="B8CCE4"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)
        # Make header text bold and centered
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                
    # Apply borders to all cells in the table
    border_attrs = {"sz": 4, "val": "single", "color": "#000000"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top=border_attrs,
                bottom=border_attrs,
                left=border_attrs,
                right=border_attrs,
            )

# ----------------------------------------
# 1. --- Cargar datos de Excel (soporte) ---
# ----------------------------------------
excel_path = r'D:\l10 report\Data\tickets (10).xlsx'
if not os.path.isfile(excel_path):
    raise FileNotFoundError(f"No se encontró el archivo de soporte: {excel_path}")

df = pd.read_excel(excel_path)

# 1.1 Normalizar 'Status'
df['Status_norm'] = df['Status'].astype(str).str.strip().str.title()

# 1.2 Normalizar 'Category'
df['Category_norm'] = (
    df['Category']
    .astype(str)
    .str.replace(r'\s*\(.*\)', '', regex=True)
    .str.strip()
    .str.title()
)

# 1.3 Totales y distribuciones (soporte)
total_tickets = len(df)

# 1.3.1 Distribución por Status
status_counts = (
    df['Status_norm']
    .value_counts()
    .reset_index()
    .rename(columns={'index': 'Status', 'Status_norm': 'Count'})
)
status_counts.columns = ['Status', 'Count']

# Calculate percentage of each status over the total tickets
total_tickets = len(df)
status_counts['Count'] = pd.to_numeric(status_counts['Count'], errors='coerce').fillna(0).astype(int)
status_counts['Percentage'] = (status_counts['Count'] / total_tickets * 100).round(2)

# Debug: print status_counts to console
print("status_counts DataFrame:")
print(status_counts.to_string(index=False))

# 1.3.2 Tickets con más de 10 días abiertos
# Define statuses to filter old tickets
focus_statuses = ["To Do", "In Progress", "QA"]
df['createdAt'] = pd.to_datetime(df['createdAt'])
df['createdAt'] = df['createdAt'].dt.tz_localize(None)
now = datetime.now()
df['DaysOpen'] = (now - df['createdAt']).dt.days

# Filter for tickets with more than 10 days open AND status 'In Progress'
df_old = df[(df['DaysOpen'] > 10) & (df['Status_norm'] == 'In Progress')].copy()
df_old = df_old.sort_values(by='DaysOpen', ascending=False)

# 1.3.3 Preparar gráfico "Distribución by Category"
category_counts = (
    df['Category_norm']
    .value_counts()
    .reset_index()
    .rename(columns={'index': 'Category', 'Category_norm': 'Count'})
)
category_counts.columns = ['Category', 'Count']
category_counts['Percentage'] = (category_counts['Count'] / total_tickets * 100).round(2)

plt.figure(figsize=(8, 6))
top20 = category_counts.head(20)
y_pos = range(len(top20))
plt.barh(y_pos, top20['Count'], align='center')
plt.yticks(y_pos, top20['Category'])
plt.xlabel('Number of Tickets')
plt.title('Distribution by Category (Top 20)')
for i, v in enumerate(top20['Count']):
    plt.text(v + 3, i, str(v), va='center')
plt.gca().invert_yaxis()
plt.tight_layout()

buf = io.BytesIO()
plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
buf.seek(0)
plt.close()

# ----------------------------------------
# 2. --- Cargar datos de JIRA ---
# ----------------------------------------
jira_path = r'D:\l10 report\Data\Jira Export CSV (all fields) 20250604054311.csv'
if not os.path.isfile(jira_path):
    raise FileNotFoundError(f"No se encontró el archivo de Jira: {jira_path}")

df_jira = pd.read_csv(jira_path, low_memory=False)
df_jira['Created'] = pd.to_datetime(df_jira['Created'], dayfirst=False, utc=True)
df_jira['Created'] = df_jira['Created'].dt.tz_localize(None)
df_jira['DaysOpen'] = (now - df_jira['Created']).dt.days

# Define the statuses we want to focus on
focus_statuses = ["To Do", "In Progress", "QA"]

# 2.1 Filtrar "USA Scaled Tickets" (Labels == 'COLSupport')
df_usa = df_jira[df_jira['Labels'].astype(str).str.contains('COLSupport', na=False)].copy()

# Filter only for the specified statuses
df_usa_filtered = df_usa[df_usa['Status'].isin(focus_statuses)].copy()

# Update the total tickets count to show only these statuses
usa_total = len(df_usa_filtered)

# Distribución USA por Status (para el gráfico), solo incluyendo los estados relevantes
usa_status_counts = (
    df_usa_filtered['Status']
    .value_counts()
    .reset_index()
    .rename(columns={'index': 'Status', 'Status': 'Count'})
)
usa_status_counts.columns = ['Status', 'Count']

# Crear gráfico horizontal de barras para USA Scaled Tickets - Status
plt.figure(figsize=(8, 4))
# Usar sort_values para ordenar los estados (To Do, In Progress, QA)
usa_status_sorted = usa_status_counts.sort_values(by='Count', ascending=True)
y_pos = range(len(usa_status_sorted))

plt.barh(y_pos, usa_status_sorted['Count'], align='center', color='#116693')  # Color azul similar al de la imagen
plt.yticks(y_pos, usa_status_sorted['Status'])
plt.xlabel('Number of Tickets')
plt.title('STATUS')
plt.grid(True, axis='x', linestyle='--', alpha=0.7)

# Añadir el conteo al final de cada barra
for i, v in enumerate(usa_status_sorted['Count']):
    plt.text(v + 0.1, i, str(v), va='center')

plt.tight_layout()

# Guardar gráfico en memoria para insertarlo luego en Word
usa_status_buf = io.BytesIO()
plt.savefig(usa_status_buf, format='png', dpi=300, bbox_inches='tight')
usa_status_buf.seek(0)
plt.close()

# Distribución USA por Priority, filtrando por los estados específicos
df_usa_status_filtered = df_usa_filtered.copy()  # Already filtered for focus_statuses

usa_priority_counts = (
    df_usa_status_filtered['Priority']  # Use the filtered dataframe 
    .value_counts()
    .reset_index()
    .rename(columns={'index': 'Priority', 'Priority': 'Count'})
)
usa_priority_counts.columns = ['Priority', 'Count']

# Crear gráfico horizontal de barras para USA Scaled Tickets - Priority
plt.figure(figsize=(8, 4))
# Ordenar por orden alfabético (Low, Medium, High) o como prefieras
usa_priority_sorted = usa_priority_counts.sort_values(by='Count', ascending=True)
y_pos = range(len(usa_priority_sorted))

plt.barh(y_pos, usa_priority_sorted['Count'], align='center', color='#116693')  # Color azul similar al anterior
plt.yticks(y_pos, usa_priority_sorted['Priority'])
plt.xlabel('Number of Tickets')
plt.title('PRIORITY')
plt.grid(True, axis='x', linestyle='--', alpha=0.7)

# Añadir el conteo al final de cada barra
for i, v in enumerate(usa_priority_sorted['Count']):
    plt.text(v + 0.1, i, str(v), va='center')

plt.tight_layout()

# Guardar gráfico en memoria para insertarlo luego en Word
usa_priority_buf = io.BytesIO()
plt.savefig(usa_priority_buf, format='png', dpi=300, bbox_inches='tight')
usa_priority_buf.seek(0)
plt.close()

usa_avg_days = df_usa_filtered.groupby('Priority')['DaysOpen'].mean().round(2)
if not usa_avg_days.empty:
    highest_priority_usa = usa_avg_days.idxmax()
    highest_avg_usa = int(usa_avg_days.max())
else:
    highest_priority_usa = None
    highest_avg_usa = 0

if highest_priority_usa:
    df_usa_top = df_usa_filtered[df_usa_filtered['Priority'] == highest_priority_usa].copy()
else:
    df_usa_top = pd.DataFrame(columns=df_usa_filtered.columns)

# 2.2 "Global Scaled Tickets" (todos los tickets de JIRA)
df_global = df_jira.copy()

# Filter global tickets by the specified statuses
df_global_filtered = df_global[df_global['Status'].isin(focus_statuses)].copy()

# Update the total for focused statuses only
global_total = len(df_global_filtered)

# Distribución Global por Status (only focused statuses)
global_status_counts = (
    df_global_filtered['Status']
    .value_counts()
    .reset_index()
    .rename(columns={'index': 'Status', 'Status': 'Count'})
)
global_status_counts.columns = ['Status', 'Count']

# Distribución Global por Priority (only for tickets with focused statuses)
global_priority_counts = (
    df_global_filtered['Priority']
    .value_counts()
    .reset_index()
    .rename(columns={'index': 'Priority', 'Priority': 'Count'})
)
global_priority_counts.columns = ['Priority', 'Count']

# "Global Tickets with Highest Priority" (filtered for focus statuses)
df_global_top = df_global_filtered[df_global_filtered['Priority'] == 'Highest'].copy()

# ----------------------------------------
# 3. --- Create a clean document from scratch ---
# ----------------------------------------
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Define status colors for consistency
status_colors = {
    'Done': 'A9D08E',       # Verde para Done
    'In Progress': 'FFD966', # Amarillo para In Progress
    'Scaled': 'BDD7EE',     # Azul claro para Scaled
    "Won't do": 'D9D9D9',   # Gris para Won't do
}

# Define the set_cell_border function used for table styling
def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#000000"},
        bottom={"sz": 12, "val": "single", "color": "#000000"},
        left={"sz": 12, "val": "single", "color": "#000000"},
        right={"sz": 12, "val": "single", "color": "#000000"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for side, attrs in kwargs.items():
        tag = f"w:{side}"
        # Create element
        elm = OxmlElement(tag)
        tcPr.append(elm)
            
        # Set attributes
        for k, v in attrs.items():
            elm.set(k, str(v))

# Create a completely new document instead of using the template
doc = Document()

# Set document properties
doc.core_properties.title = "Support & Tickets Report"
doc.core_properties.author = "L10 Team"

# Create title
title = doc.add_heading('Support & Tickets Report', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(24)

# ----------------------------------------
# NEW SECTION: Recent Support Tickets (Past Two Weeks)
# ----------------------------------------
# Filter tickets from the past two weeks
two_weeks_ago = now - timedelta(days=14)
df_recent = df[df['createdAt'] >= two_weeks_ago].copy()
recent_tickets_count = len(df_recent)

# ----------------------------------------
# Add calls data processing
# ----------------------------------------
# Load and process calls data
calls_path = r'D:\l10 report\Data\Calls.csv'
if os.path.isfile(calls_path):
    df_calls = pd.read_csv(calls_path)
    # Convert Unix timestamp to datetime
    df_calls['DateTime'] = pd.to_datetime(df_calls['Time'], unit='s')
    
    # Filter calls from the past two weeks
    recent_calls = df_calls[df_calls['DateTime'] >= two_weeks_ago]
    recent_calls_count = len(recent_calls)
else:
    recent_calls_count = 0

# Add section header
section = doc.add_heading('Recent Support Activity (Past Two Weeks)', level=2)
section.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in section.runs:
    run.font.size = Pt(18)

# Create a table for total counts
total_counts_table = doc.add_table(rows=3, cols=2)
total_counts_table.style = 'Table Grid'

# Add headers
header_cells = total_counts_table.rows[0].cells
header_cells[0].text = 'Activity Type'
header_cells[1].text = 'Total Count'

# Format header
for cell in header_cells:
    # Set light blue background (#B8CCE4) for header
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="B8CCE4"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True

# Add call count row
call_cells = total_counts_table.rows[1].cells
call_cells[0].text = 'Phone Calls'
call_cells[1].text = str(recent_calls_count)
call_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add ticket count row
ticket_cells = total_counts_table.rows[2].cells
ticket_cells[0].text = 'Support Tickets'
ticket_cells[1].text = str(recent_tickets_count)
ticket_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Apply borders to all cells in the table
for row in total_counts_table.rows:
    for cell in row.cells:
        set_cell_border(
            cell,
            top={"sz": 4, "val": "single", "color": "#000000"},
            bottom={"sz": 4, "val": "single", "color": "#000000"},
            left={"sz": 4, "val": "single", "color": "#000000"},
            right={"sz": 4, "val": "single", "color": "#000000"},
        )

# Add some space after the table
doc.add_paragraph()

# Create pie chart of status distribution for recent tickets
if recent_tickets_count > 0:
    # Changed from Status_norm to Category_norm to show categories instead of statuses
    recent_category_counts = (
        df_recent['Category_norm']
        .value_counts()
        .reset_index()
    )
    # Explicitly rename the columns to ensure correct names
    recent_category_counts.columns = ['Category', 'Count']
    
    # Calculate percentages for each category
    recent_category_counts['Percentage'] = (recent_category_counts['Count'] / recent_category_counts['Count'].sum() * 100).round(2)
    
    # Group categories with less than 3% into "Other"
    threshold = 3.0
    main_categories = recent_category_counts[recent_category_counts['Percentage'] >= threshold]
    small_categories = recent_category_counts[recent_category_counts['Percentage'] < threshold]
    
    if not small_categories.empty:
        # Create "Other" row with sum of all small categories
        other_row = pd.DataFrame({
            'Category': ['Other'],
            'Count': [small_categories['Count'].sum()],
            'Percentage': [small_categories['Percentage'].sum()]
        })
        
        # Combine main categories with "Other" category
        plot_data = pd.concat([main_categories, other_row], ignore_index=True)
    else:
        plot_data = main_categories
    
    # Create pie chart
    plt.figure(figsize=(7, 5))
    colors = ['#A9D08E', '#FFD966', '#BDD7EE', '#D9D9D9', '#F4B084', '#C6E0B4', '#FFE699', '#9BC2E6', '#C9C9C9', '#F8CBAD']  # Using more colors for categories
    
    # If there are more categories than colors, it will cycle through the colors
    plt.pie(
        plot_data['Count'], 
        labels=plot_data['Category'], 
        autopct='%1.1f%%',
        startangle=90,
        colors=colors[:len(plot_data)]
    )
    plt.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle
    plt.title('Recent Tickets by Category')  # Changed title to reflect categories
    
    # Save pie chart to buffer
    recent_category_buf = io.BytesIO()
    plt.savefig(recent_category_buf, format='png', dpi=300, bbox_inches='tight')
    recent_category_buf.seek(0)
    plt.close()
    
    # Add pie chart to document
    chart_para = doc.add_paragraph()
    chart_run = chart_para.add_run()
    chart_run.add_picture(recent_category_buf, width=Inches(5))  # Using the category buffer
    
    
# Add total tickets
doc.add_heading('Total support tickets', level=3)
section.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in section.runs:
    run.font.size = Pt(16)
doc.add_paragraph(f'Total Tickets: {total_tickets}')

# Create the STATUS table
status_table = doc.add_table(rows=1, cols=3)
status_table.style = 'Table Grid'

# Add header row with dark gray background and white text
header_cells = status_table.rows[0].cells
header_cells[0].text = 'STATUS'
header_cells[1].text = '%'
header_cells[2].text = '# TICKETS'

# Format header
for i, cell in enumerate(header_cells):
    # Set dark gray background
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="404040"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Set white text and center alignment
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(11)

# Add data rows with colors
for _, row_data in status_counts.iterrows():
    row_cells = status_table.add_row().cells
    status = row_data['Status']
    row_cells[0].text = status
    row_cells[1].text = f"{row_data['Percentage']}%"
    row_cells[2].text = str(int(row_data['Count']))
    
    # Apply colors
    color = status_colors.get(status, 'FFFFFF')
    for i, cell in enumerate(row_cells):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Center content in columns 1 and 2
        if i > 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Set grid borders for the table
for row in status_table.rows:
    for cell in row.cells:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#000000"},
            bottom={"sz": 12, "val": "single", "color": "#000000"},
            left={"sz": 12, "val": "single", "color": "#000000"},
            right={"sz": 12, "val": "single", "color": "#000000"},
        )

# Add Category Distribution section
section = doc.add_heading('Distribution by Category:', level=1)
section.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in section.runs:
    run.font.size = Pt(18)
category_para = doc.add_paragraph()
category_run = category_para.add_run()
category_run.add_picture(buf, width=Inches(6))

# Add Tickets with más de 10 días abiertos
doc.add_paragraph()
doc.add_paragraph()
section = doc.add_heading('Tickets with more than 10 days open', level=2)

section.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in section.runs:
    run.font.size = Pt(18)
old_tickets_table = doc.add_table(rows=1, cols=5)
old_tickets_table.style = 'Table Grid'

# Add header
old_header = old_tickets_table.rows[0].cells
old_header[0].text = 'ID Ticket'
old_header[1].text = 'Company'
old_header[2].text = 'Contact'  # Changed from Reporter to Contact
old_header[3].text = 'Description'
old_header[4].text = 'Days Open'

# Format header
for cell in old_header:
    cell.paragraphs[0].runs[0].font.bold = True

# Add data rows
for _, row in df_old.iterrows():
    cells = old_tickets_table.add_row().cells
    cells[0].text = str(int(row['IDTicket']))
    cells[1].text = str(row['Companyname'])
    cells[2].text = str(row.get('Contact', ''))  # Using Contact column instead of Reporter
    cells[3].text = str(row['Description'])
    cells[4].text = str(int(row['DaysOpen']))

# Add ticket IDs
for ticket_id in df_old['IDTicket']:
    doc.add_paragraph(f"{int(ticket_id)}:")

# ----------------------------------------
# 4. --- USA Scaled Tickets Section ---
# ----------------------------------------
doc.add_paragraph()
doc.add_paragraph()
section = doc.add_heading('USA Scaled Tickets', level=2)

section.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in section.runs:
    run.font.size = Pt(18)
doc.add_paragraph(f'Total tickets: {usa_total}')

# Add STATUS graph
status_para = doc.add_paragraph()
status_run = status_para.add_run()
status_run.add_picture(usa_status_buf, width=Inches(6))

# Add Priority section
doc.add_paragraph('Priority')
# Add priority graph
priority_para = doc.add_paragraph()
priority_run = priority_para.add_run()
priority_run.add_picture(usa_priority_buf, width=Inches(6))


# Add Highest average days
doc.add_paragraph(f'Highest average days opened = {highest_avg_usa} days')

# Add USA Top Priority Tickets
section = doc.add_heading('USA Tickets with Top Priority', level=2)
section.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in section.runs:
    run.font.size = Pt(16)
usa_tickets_table = doc.add_table(rows=1, cols=7)
hdr = usa_tickets_table.rows[0].cells
hdr[0].text = 'Issue key'
hdr[1].text = 'Issue id'
hdr[2].text = 'Summary'
hdr[3].text = 'Priority'
hdr[4].text = 'Status'
hdr[5].text = 'Created'
hdr[6].text = 'Sprint'

style_table_like_image(usa_tickets_table)
set_column_width(usa_tickets_table, 2, 3.0)

# Filter for Highest and High priority tickets in focus statuses only
df_usa_highest = df_usa_filtered[df_usa_filtered['Priority'] == 'Highest'].copy().sort_values(by='Created', ascending=False)
df_usa_high = df_usa_filtered[df_usa_filtered['Priority'] == 'High'].copy().sort_values(by='Created', ascending=False)

# First add all Highest priority tickets
added_rows = 0
for _, row in df_usa_highest.head(10).iterrows():
    r = usa_tickets_table.add_row().cells
    r[0].text = str(row['Issue key'])
    r[1].text = str(row.get('Issue id', ''))
    r[2].text = str(row.get('Summary', ''))
    r[3].text = str(row.get('Priority', ''))
    r[4].text = str(row.get('Status', ''))
    
    created_date = row.get('Created')
    if created_date is not None and hasattr(created_date, 'strftime'):
        r[5].text = created_date.strftime('%m/%d/%Y')
    else:
        r[5].text = str(created_date)
        
    r[6].text = str(row.get('Sprint', ''))
    added_rows += 1

# If we have fewer than 5 highest priority tickets, add some high priority ones too
if added_rows < 5:
    for _, row in df_usa_high.head(10 - added_rows).iterrows():
        r = usa_tickets_table.add_row().cells
        r[0].text = str(row['Issue key'])
        r[1].text = str(row.get('Issue id', ''))
        r[2].text = str(row.get('Summary', ''))
        r[3].text = str(row.get('Priority', ''))
        r[4].text = str(row.get('Status', ''))
        
        created_date = row.get('Created')
        if created_date is not None and hasattr(created_date, 'strftime'):
            r[5].text = created_date.strftime('%m/%d/%Y')
        else:
            r[5].text = str(created_date)
            
        r[6].text = str(row.get('Sprint', ''))

# ----------------------------------------
# 5. --- Global Scaled Tickets Section ---
# ----------------------------------------
doc.add_paragraph()
doc.add_paragraph()
section = doc.add_heading('Global Scaled Tickets', level=3)
section.alignment = WD_ALIGN_PARAGRAPH.CENTER
section.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in section.runs:
    run.font.size = Pt(18)
doc.add_paragraph(f'Total Tickets: {global_total}')

# Add Status section
doc.add_paragraph('Status')
global_status_table = doc.add_table(rows=1, cols=2)
hdr = global_status_table.rows[0].cells
hdr[0].text = 'Status'
hdr[1].text = 'Count of Status'

# Format table
style_table_like_image(global_status_table)

# Add data rows
for _, row in global_status_counts.iterrows():
    r = global_status_table.add_row().cells
    r[0].text = str(row['Status'])
    r[1].text = str(int(row['Count']))

# Add Priority section
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph('Priority')
global_priority_table = doc.add_table(rows=1, cols=2)
hdr = global_priority_table.rows[0].cells
hdr[0].text = 'Priority'
hdr[1].text = 'Count of Priority'

# Format table
style_table_like_image(global_priority_table)

# Add data rows
for _, row in global_priority_counts.iterrows():
    r = global_priority_table.add_row().cells
    r[0].text = str(row['Priority'])
    r[1].text = str(int(row['Count']))

# Add Global Top Priority Tickets
doc.add_paragraph()
section = doc.add_heading('Global Tickets with Highest Priority', level=3)
section.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in section.runs:
    run.font.size = Pt(16)
global_tickets_table = doc.add_table(rows=1, cols=7)
hdr = global_tickets_table.rows[0].cells
hdr[0].text = 'Issue key'
hdr[1].text = 'Issue id'
hdr[2].text = 'Summary'
hdr[3].text = 'Priority'
hdr[4].text = 'Status'
hdr[5].text = 'Created'
hdr[6].text = 'Sprint'

style_table_like_image(global_tickets_table)
set_column_width(global_tickets_table, 2, 3.0)

# Add data rows - sort by creation date (newest first)
sorted_global_top = df_global_top.sort_values(by='Created', ascending=False)

for _, row in sorted_global_top.head(10).iterrows():
    r = global_tickets_table.add_row().cells
    r[0].text = str(row['Issue key'])
    r[1].text = str(row.get('Issue id', ''))
    r[2].text = str(row.get('Summary', ''))
    r[3].text = str(row.get('Priority', ''))
    r[4].text = str(row.get('Status', ''))
    
    created_date = row.get('Created')
    if created_date is not None and hasattr(created_date, 'strftime'):
        r[5].text = created_date.strftime('%m/%d/%Y %H:%M')
    else:
        r[5].text = str(created_date)
        
    r[6].text = str(row.get('Sprint', ''))

# ----------------------------------------
# 6. --- Guardar documento final con fecha en el nombre ---
# ----------------------------------------
# Close any previous handles to the file that might be causing corruption
import gc
gc.collect()  # Force garbage collection to release any memory/file handles

fecha = now.strftime('%Y-%m-%d')
hora = now.strftime('%H-%M-%S')
output_dir = r'D:\l10 report\reports'
os.makedirs(output_dir, exist_ok=True)
output_filename = os.path.join(output_dir, f'test_de_reporte_l10_full_{fecha}_time_{hora}.docx')

# Save with additional error handling
try:
    doc.save(output_filename)
    print(f'Reporte generado: {output_filename}')
except Exception as e:
    print(f'Error al guardar el documento: {e}')
    # Try a different approach to save
    try:
        # Create a new document and copy content
        from copy import deepcopy
        temp_doc = Document()
        
        # Copy styles if possible
        if hasattr(doc, 'styles') and hasattr(temp_doc, 'styles'):
            for style in doc.styles:
                if style.name not in temp_doc.styles:
                    try:
                        temp_doc.styles.add_style(style.name, style.style_id)
                    except:
                        pass
        
        # Copy content - paragraphs
        for para in doc.paragraphs:
            p = temp_doc.add_paragraph()
            for run in para.runs:
                r = p.add_run(run.text)
                # Copy run formatting
                r.bold = run.bold
                r.italic = run.italic
                r.underline = run.underline
                if hasattr(run, 'font') and hasattr(r, 'font'):
                    if hasattr(run.font, 'name'):
                        r.font.name = run.font.name
                    if hasattr(run.font, 'size'):
                        r.font.size = run.font.size
        
        # Copy tables (simple copy)
        for table in doc.tables:
            t = temp_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            # Copy cell content
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if i < len(t.rows) and j < len(t.rows[i].cells):
                        t.rows[i].cells[j].text = cell.text
        
        # Save the new document
        temp_filename = os.path.join(output_dir, f'fixed_report_{fecha}_time_{hora}.docx')
        temp_doc.save(temp_filename)
        print(f'Reporte recuperado generado: {temp_filename}')
    except Exception as e2:
        print(f'Error al intentar recuperar documento: {e2}')